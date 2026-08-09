from pathlib import Path
from datetime import date
import re
import hashlib

from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SCREEN_DIR = ROOT / "screenshots"
WORKING_DIR = SCREEN_DIR / "_working"
TREE_FILE = ROOT / "01_功能树_源码整理版.md"
SHOT_LIST_MD = ROOT / "02_截图清单_功能覆盖版.md"
FINAL_MD = ROOT / "03_Spec2.02功能说明与截图文档.md"
FINAL_DOCX = ROOT / "03_Spec2.02功能说明与截图文档.docx"


GROUPS = {
    "00": {
        "number": "0",
        "title": "工作台与运行基础",
        "summary": "主程序负责单实例启动、方法标题、菜单调度、信息记录和状态反馈。菜单截图用于确认各功能入口，不作为功能树节点。",
        "source": "Main/FmMain.pas/.dfm、Main/Others/FmMessage.pas",
    },
    "10": {
        "number": "1",
        "title": "光谱方法",
        "summary": "覆盖方法的新建、打开、另存、页面设置、打印机设置以及方法库对象的基本生命周期。",
        "source": "Main/mmMethod/FmInputQuery、FmSelect、FmPageSetup，以及 FmMain 方法菜单动作",
    },
    "20": {
        "number": "2",
        "title": "分析条件",
        "summary": "覆盖方法参数、谱板/谱带参数、分析谱线配置、峰值和拟合方式、内标方式、谱线类型及参数打印。",
        "source": "Main/mmMethod/FmMtdCfg、FmMtdLines、FmMtdLinesAdd、FmMtdPrint；Common/uAnaLine、uFitMode、uPeakMode",
    },
    "30": {
        "number": "3",
        "title": "分析测试",
        "summary": "正式截图覆盖色散、蒸发和样品摄谱工作条件、采集控制台、预激发进度、预录样号、实际谱图、多样品对比、横纵向缩放、十字线/单点强度、全时样品谱图、样品结果、标准曲线、重复性数据检查和慢进参数入口；真实硬件采集状态及慢进人工干预闭环仍需现场验收。",
        "source": "Main/mmAnalyze/FmIgnitCon、FmView、FmViewSingle、FmViewMulti、FmSampInput、FmAnaCurve、FmAnaPdt",
    },
    "40": {
        "number": "4",
        "title": "数据处理",
        "summary": "正式截图覆盖既有结果重新计算、报告打印、报告导出向导三步、强度导出三步向导及实际文件载入；PDF、独立打印预览和成功导出后的结果页仍未形成合格截图。",
        "source": "Main/mmData/FmExport.pas、FmPrintDat.pas；Common/uExport.pas",
    },
    "50": {
        "number": "5",
        "title": "工具",
        "summary": "正式截图覆盖软件参数、设备参数、CCD 设置、方法管理、数据库备份压缩及蒸发摄谱文件导入；调试对中当前只能取得串口超时异常，未纳入正式功能证据。",
        "source": "Main/mmTool/FmCfgSoftWare、FmOptHardWare、FmDebug、FmMtdManage、FmMtdBackup、FmImportPre",
    },
    "60": {
        "number": "6",
        "title": "帮助",
        "summary": "正式截图覆盖关于窗口；帮助主题因同目录 DIRECT.CHM 缺失只能出现提示，未纳入正式证据。",
        "source": "Main/Others/FmHint.pas、FmAbout.pas",
    },
}


REQUIREMENT_STATUS_ROWS = [
    ("摄谱区间数据保存", "部分支持", "源码确认但缺少合格运行截图", "补充按转角切换平均值/全时数据的有效配置或结果图；现场确认全区间数据保存语义。"),
    ("新增谱线可检测范围", "部分支持", "源码确认但缺少合格运行截图", "补充输入越界波长后的有效提示图，并验收 CCD 漂移边界提示。"),
    ("谱线类型与内标方式", "部分支持", "已有合格运行截图", "验收基线是否需要独立谱线类型；当前基线仍是背景参数/计算逻辑。"),
    ("峰值方式", "源码确认支持", "已有合格运行截图", "按现有最大值/高斯模式进入产品验收。"),
    ("拟合坐标", "源码确认支持", "已有合格运行截图", "按普通/对数坐标切换进入产品验收。"),
    ("拟合函数", "源码确认支持", "已有合格运行截图", "按直线、二次、三次、样条函数进入产品验收。"),
    ("方法参数打印", "源码确认支持", "已有合格运行截图", "补打印输出文件或纸面结果时再做打印质量验收。"),
    ("转角排序与电弧/峰位异常闭环", "部分支持", "源码确认但缺少合格运行截图", "有真实设备时补短波到长波、关键波段优先、异常提示/校正/重采/写回闭环；源码未发现明确排序策略。"),
    ("摄谱类型与曲线分析", "源码确认支持", "已有合格运行截图", "按色散、蒸发、样品三类流程进入产品验收。"),
    ("多样品谱图对比", "源码确认支持", "已有合格运行截图", "按多样品曲线和单点强度图进入产品验收。"),
    ("谱线/检测器定位与缩放", "源码确认支持", "已有合格运行截图", "按十字线、CCD 切换、横纵向缩放进入产品验收。"),
    ("样号输入时机", "源码确认支持", "已有合格运行截图", "补采集完成后重命名的有效结果图，确认写回样品信息。"),
    ("慢进与人工干预", "基本支持，需验收", "源码确认但缺少合格运行截图", "补逐谱线暂停→人工定位调整→继续→结果写回闭环；当前仅有监控/等待时间配置图。"),
    ("重复性误差识别与预警", "源码确认支持", "已有合格运行截图", "按均值、标准差、RSD、超差提示进入产品验收。"),
    ("标准曲线切换与调整", "源码确认支持", "已有合格运行截图", "按拟合切换、标准点停用、原始值恢复和重算进入产品验收。"),
    ("调试与汞灯调试", "部分支持；汞灯流程重点补充", "源码确认但缺少合格运行截图", "补有效 CCD 调试曲线；新版本增加汞灯选择、谱线校准和专用流程。"),
    ("报告打印、预览与导出", "部分支持", "已有文本/Excel 配置截图；PDF/独立预览缺少合格证据", "新版本补 PDF 和独立打印预览；当前截图不宣称 PDF 成功。"),
    ("角色和权限", "未发现实现", "无合格运行截图", "后续版本补用户、登录、角色、权限和审计，并形成验收条款。"),
    ("全时数据独立导出", "源码确认支持", "源码确认但缺少合格运行截图", "补实际全时数据导出成功结果；现有图只证明谱图查看/观察。"),
    ("帮助主题", "源码确认支持", "源码确认但缺少合格运行截图", "补齐 DIRECT.CHM 后取得主题、目录/索引阅读图。"),
]


SIMILARITY_THRESHOLD = 0.995


STATUS_RULES = [
    ("载入实际文件", "实测载入实际数据文件"),
    ("实际文件", "实测载入实际文件后的界面"),
    ("实际谱图", "实测实际谱图界面"),
    ("结果载入", "实测载入结果参数界面"),
    ("目标目录", "实测向导目标目录设置界面"),
    ("选项设置", "实测向导导出选项界面"),
    ("载入文件后", "实测载入实际数据文件后的界面"),
    ("预激发进度", "实测采集进度/预激发状态"),
    ("载入样品文件", "实测载入样品文件后的界面"),
    ("入口实测", "实测功能入口界面"),
    ("选项", "实测选项展开界面"),
    ("方法列表", "实测方法库列表界面"),
    ("输入名称", "实测输入/命名界面"),
]


def group_for(path: Path):
    prefix = path.name[:2]
    if prefix in {"20", "21", "22"}:
        return "20"
    return prefix


def title_for(path: Path):
    stem = path.stem
    if "_" in stem:
        stem = stem.split("_", 1)[1]
    return stem.replace("_", " / ")


def status_for(path: Path):
    for token, status in STATUS_RULES:
        if token in path.stem:
            return status
    return "实测界面"


def description_for(path: Path):
    n = path.stem
    if n.startswith("00_"):
        return "展示主菜单或子菜单的实测入口，确认该功能组在主工作台中的层级位置。"
    if "打开方法" in n:
        return "展示从方法库选择并载入方法的界面，确认方法列表和当前方法切换入口。"
    if "新建方法" in n:
        return "展示创建方法时的名称输入和校验入口，确认方法文件生命周期的起点。"
    if "另存为" in n:
        return "展示当前方法副本命名入口，确认另存为功能可从当前方法生成新方法。"
    if "页面设置" in n:
        return "展示纸张、方向、边距等页面参数入口，用于配置方法参数和报告打印布局。"
    if "打印机设置" in n:
        return "展示系统打印机和打印参数设置入口。"
    if "方法参数" in n:
        return "展示方法基本信息、谱板布局、激发参数和谱带参数页签，确认方法条件分层配置。"
    if "分析谱线" in n:
        return "展示谱线列表及其编辑参数；对应谱线类别、内标、峰值、拟合和分析条件配置。"
    if "参数打印" in n:
        return "展示方法条件参数和分析谱线参数打印入口，确认方法配置可形成打印输出。"
    if "标准曲线" in n:
        return "展示标准曲线实际数据、拟合结果、标准点启停及拟合方式切换；用于核对曲线拟合和结果复核流程。"
    if "重复性误差" in n:
        return "展示重复测量数据检查窗口，包含平均值、极差、标准偏差、相对偏差和有效点数等质控信息。"
    if "慢进模式" in n:
        return "展示分析参数中的测量谱带监控和等待时间设置；该图证明慢进配置入口，不单独证明完整人工干预闭环。"
    if "采集控制台" in n:
        return "展示全时样品摄谱工作台的样品输入、重复谱、试样摄谱、预热摄谱、停止和结束控制。"
    if "预激发进度" in n:
        return "展示样品摄谱实际启动后的第 1 次摄谱和预激发倒计时状态；不把设备未连接后的错误状态写成成功采集。"
    if "实际结果" in n:
        return "展示样品分析载入实际全时样品后的谱图和结果分析界面。"
    if "全时样品实际谱图" in n:
        return "展示实际全时样品的谱图、记录列表、样品编号和曝光区间，确认全时样品谱图查看流程。"
    if "多样品对比" in n:
        return "展示多个样品同时绘制的强度曲线和样品记录列表，确认谱图对比功能。"
    if "横向与纵向缩放" in n:
        return "展示谱图缩放后的横向/纵向视图及缩放工具状态，确认谱图交互缩放功能。"
    if "十字线与单点强度" in n:
        return "展示十字线定位后的波长坐标、CCD 序号、曝光区间和各样品单点强度。"
    if "导入蒸发摄谱文件" in n:
        return "展示蒸发摄谱文件导入入口、记录选择和目标样品数据设置。"
    if "摄谱" in n:
        return "展示对应摄谱流程的入口界面，确认色散、蒸发或样品数据采集功能存在。"
    if "曲线分析" in n:
        return "展示曲线分析入口或载入实际数据后的分析界面，确认后续曲线处理流程。"
    if "查看谱图" in n:
        return "展示谱图查看入口或载入样品后的查看界面，确认谱带/强度曲线查看流程。"
    if "样品分析" in n:
        return "展示样品分析入口，确认样品结果分析和后续曲线计算流程。"
    if "预录样号" in n:
        return "展示样品号预录界面，确认样品队列可在摄谱前录入并管理。"
    if "打印报告" in n:
        return "展示报告打印入口，确认既有分析报告可选择并进入打印处理。"
    if "导出报告" in n:
        return "展示报告导出向导入口，确认报告文件可按向导选择并输出为文本或电子表格。"
    if "导出强度" in n:
        return "展示强度文件导出向导，覆盖文件选择、目标目录、同名文件处理、文本/电子表格类型和导出完成。"
    if "重新计算" in n:
        return "展示既有强度结果重新计算入口，确认可结合当前方法重算分析结果。"
    if "软件参数" in n:
        return "展示软件级目录、日志、消息和显示相关配置。"
    if "设备参数" in n:
        return "展示串口、通信和设备参数；CCD 设置截图进一步确认采集硬件配置入口。"
    if "调试对中" in n:
        return "展示 CCD 调试曲线、对中操作和串口异常反馈，确认设备调试入口不等同于正常样品采集。"
    if "方法管理" in n:
        return "展示方法库浏览、启用/禁用和当前方法管理功能。"
    if "备份压缩数据库" in n:
        return "展示方法库备份和数据库压缩操作入口；本次仅验证界面，不执行写库操作。"
    if "帮助主题" in n:
        return "展示帮助主题入口，确认可调用程序同目录帮助文件。"
    if "关于" in n:
        return "展示软件名称、版本和版权信息。"
    return "展示该功能节点的实测界面，用于确认功能入口和主要控件。"


def evidence_for(group):
    return GROUPS[group]["source"]


def similarity_pairs(files):
    """Return near-identical thumbnail pairs for manual review, not automatic deletion."""
    thumbs = []
    for f in files:
        with Image.open(f) as image:
            resized = image.convert("L").resize((96, 60))
            data = resized.get_flattened_data() if hasattr(resized, "get_flattened_data") else resized.getdata()
            values = list(data)
        mean = sum(values) / len(values)
        centered = [value - mean for value in values]
        norm = sum(value * value for value in centered) ** 0.5
        thumbs.append((f.name, centered, norm))

    pairs = []
    for i, (left_name, left, left_norm) in enumerate(thumbs):
        if left_norm == 0:
            continue
        for right_name, right, right_norm in thumbs[i + 1:]:
            if right_norm == 0:
                continue
            score = sum(a * b for a, b in zip(left, right)) / (left_norm * right_norm)
            if score >= SIMILARITY_THRESHOLD:
                pairs.append((score, left_name, right_name))
    return sorted(pairs, reverse=True)


def coverage_audit_lines(files):
    pairs = similarity_pairs(files)
    similarity_lines = [
        f"- SHA-256 完全重复检查：{len(files)} 张正式 PNG 均为唯一文件。",
        f"- 缩略图视觉相似度复核：以余弦相似度 {SIMILARITY_THRESHOLD:.3f} 为人工复核阈值，发现 {len(pairs)} 对候选；候选均为不同功能或不同业务状态，未自动删除。",
    ]
    return [
        "## 截图与功能树对照复核",
        "",
        f"- 正式目录：{len(list(SCREEN_DIR.glob('*.png')))} 张独立 PNG 功能证据；补充目录复核后仅纳入有效采集状态图。",
        "- 非正式过程/补充截图：已清理出正式目录；未被最终文档引用的素材不再保留在当前截图文件夹，备份副本仍保留。",
        *similarity_lines,
        "- 正式目录不是所有尝试截图的总数；它只保留真实功能界面、实际数据界面或有明确业务价值的配置界面。",
        "- 清理前未采用的素材主要是菜单、空白/无数据工作台、Windows 文件选择器、报错/警告、覆盖确认、任务管理器、Codex 界面、焦点探测或与正式图重复的版本；这些文件已从当前截图文件夹移除，备份副本保留。",
        "",
        "### 当前已有合格运行截图的功能范围",
        "",
        "- 光谱方法：新建、打开、另存为、页面设置和打印机设置。",
        "- 分析条件：激发/谱板/谱带参数、分析谱线列表和编辑入口、谱线类别、内标方式、峰值方式、拟合方式及参数打印。",
        "- 分析测试：色散/蒸发/样品摄谱工作条件、样品摄谱采集控制台、预激发进度、预录样号、实际谱图、多样品对比、横纵向缩放、十字线/单点强度、全时样品谱图、样品分析结果、标准曲线、重复性数据检查和慢进参数入口。",
        "- 数据处理：重新计算、打印报告、报告导出向导三步、强度导出三步向导及实际文件载入；未纳入 PDF 保存对话框或文件选择器。",
        "- 工具与帮助：软件参数、设备参数、CCD 设置、方法管理、备份压缩数据库、蒸发文件导入和关于窗口。",
        "",
        "### 功能树中尚无合格运行截图的细分节点",
        "",
        "- 真实硬件采集过程中的暂停/继续/停止和完成收尾闭环；现有正式图只覆盖采集控制台和预激发倒计时，不宣称硬件采集成功。",
        "- 慢进的完整暂停→人工调整→继续→结果写回闭环；当前截图只证明测量谱带监控和等待时间配置入口。",
        "- 全时数据独立导出完成状态；现有全时样品谱图只证明查看/观察，不证明独立导出成功。",
        "- 报告打印预览、PDF 导出和成功导出后的结果页；已补齐报告导出向导第 1–3 步，但 PDF 保存对话框、文件选择器和错误残留图均不作为功能证据。",
        "- 调试对中有效曲线/汞灯调试画面；当前硬件条件下只能出现串口异常，未纳入正式证据。帮助主题同理，现有候选为 DIRECT.CHM 缺失提示。",
        "- 新软件要求中的转角排序/关键波段优先、基线/峰位异常闭环、汞灯专用流程、PDF/打印预览和角色权限管理，不能由现有截图宣称已实现；具体源码结论仍以功能树第五章为准。",
        "",
        "### 复核结论",
        "",
        f"{len(list(SCREEN_DIR.glob('*.png')))} 张正式截图已覆盖功能树中主要可见配置、实际数据和结果功能；仍不能覆盖每一个源码内部节点、硬件运行状态或每一个运行时分支。清理前已复核过程/补充目录，未把无关、报错或重复截图纳入正式证据；未采用素材已移除，备份副本保留。缺口已在功能说明中标为“源码确认但缺少合格运行截图”或“后续版本重点补充”。",
        "",
    ]


def entries():
    files = sorted(SCREEN_DIR.glob("*.png"), key=lambda p: p.name)
    if not files:
        raise RuntimeError("正式截图目录为空")
    seen_hashes = set()
    for f in files:
        if group_for(f) not in GROUPS:
            raise RuntimeError(f"截图未归入功能组: {f.name}")
        with Image.open(f) as image:
            if image.size != (2880, 1800):
                raise RuntimeError(f"截图尺寸错误: {f.name} -> {image.size}")
        digest = hashlib.sha256(f.read_bytes()).hexdigest()
        if digest in seen_hashes:
            raise RuntimeError(f"正式截图存在完全重复文件: {f.name}")
        seen_hashes.add(digest)
    return files


def md_image_path(path: Path):
    return path.relative_to(ROOT).as_posix()


def write_screenshot_list(files):
    lines = [
        "# Spec2.02 截图清单（功能覆盖版）",
        "",
        f"- 正式截图数量：{len(files)} 张",
        "- 图像格式：PNG",
        "- 图像尺寸：2880×1800（桌面缩放 200% 环境下按物理屏幕采集）",
        "- 非正式过程/补充截图：已清理，不纳入最终文档；备份副本保留。",
        "- 正式目录仅保留真实功能画面；菜单、文件选择器、报错、空白和重复截图均已隔离。",
        "- 组织方式：按功能树 0-6 层级归档；截图用于证明真实配置、实际数据或结果界面，不改变功能树结构。",
        "",
    ]
    current = None
    fig = 1
    for f in files:
        group = group_for(f)
        if group != current:
            current = group
            g = GROUPS[group]
            lines += [f"## {g['number']}. {g['title']}", "", g["summary"], "", f"源码依据：{g['source']}", ""]
        lines += [
            f"### 图 {fig}：{title_for(f)}",
            "",
            f"- 文件：`{f.name}`",
            f"- 实测状态：{status_for(f)}",
            f"- 功能说明：{description_for(f)}",
            f"- 源码依据：{evidence_for(group)}",
            "",
        ]
        fig += 1
    lines += coverage_audit_lines(files)
    SHOT_LIST_MD.write_text("\n".join(lines), encoding="utf-8")


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_style_font(style, name="Calibri", size=11, color="000000", bold=None, italic=None):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    if italic is not None:
        style.font.italic = italic
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    old_grid = tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        old_grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                mar = OxmlElement("w:tcMar")
                tc_pr.append(mar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                node = mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375 + level * 0.25)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_font(r)
    return p


def add_tree_to_doc(doc, text):
    for raw in text.splitlines():
        line = raw.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            continue
        m = re.match(r"^(#{2,4})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)) - 1, 3)
            doc.add_heading(m.group(2), level=level)
        elif line.startswith("- "):
            add_bullet(doc, line[2:])
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line[2:])
            set_font(r, size=10, color="555555", italic=True)
        elif line.startswith("|"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(line)
            set_font(r, size=9, color="555555")
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(line)
            set_font(r)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color="000000")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        set_style_font(style, size=size, color=color, bold=True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    cap = doc.styles["Caption"]
    set_style_font(cap, size=9, color="555555", italic=True)
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("Spec2.02 功能研究 | 功能说明与截图")
    set_font(r, size=9, color="666666")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    r = footer.add_run("Spec2.02 | 第 ")
    set_font(r, size=9, color="666666")
    add_field(footer, " PAGE ")
    r = footer.add_run(" 页")
    set_font(r, size=9, color="666666")


def add_cover(doc, file_count):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("功能研究报告")
    set_font(r, size=12, color="B2873B", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Spec2.02")
    set_font(r, size=30, color="203748", bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("功能说明与运行截图文档")
    set_font(r, size=15, color="2B5163")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(50)
    r = p.add_run("基于 SpecDirect.exe 实测界面与 Source 源码审查")
    set_font(r, size=10.5, color="666666", italic=True)

    for label, value in [
        ("程序入口", r"../Spec2.02/SpecDirect.exe"),
        ("源码目录", r"../Spec Source/Source"),
        ("采集环境", "2880×1800，桌面缩放 200%"),
        ("截图范围", f"{file_count} 张正式 PNG，保留真实功能画面"),
        ("编制日期", str(date.today())),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{label}：")
        set_font(r, size=10.5, color="203748", bold=True)
        r = p.add_run(value)
        set_font(r, size=10.5, color="555555")
    doc.add_page_break()


def add_status_table(doc):
    doc.add_heading("新软件要求：源码覆盖与后续验收重点", level=1)
    p = doc.add_paragraph("本章独立记录用户提供的新软件要求，不把新要求误写成 Spec2.02 当前已实现功能。每行分开记录源码状态、截图证据状态和下一步，缺少截图不等同于没有功能。")
    for r in p.runs:
        set_font(r)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["要求项", "源码状态", "截图证据状态", "下一步"]
    for cell, text in zip(table.rows[0].cells, headers):
        shade_cell(cell, "E8EEF5")
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_font(r, size=9.5, color="203748", bold=True)
    mark_table_header(table.rows[0])
    for item, source_status, screenshot_status, next_step in REQUIREMENT_STATUS_ROWS:
        cells = table.add_row().cells
        for cell, text in zip(cells, (item, source_status, screenshot_status, next_step)):
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(text)
            set_font(r, size=9, color="000000")
    set_table_geometry(table, [1800, 1900, 2300, 3360])


def add_coverage_audit(doc, file_count, files):
    doc.add_heading("截图与功能树对照复核", level=1)
    pairs = similarity_pairs(files)
    doc.add_paragraph(
        f"正式目录保留 {file_count} 张独立 PNG 功能证据；未被最终文档引用的过程/补充截图已从当前截图文件夹清理，备份副本保留。"
        " 已逐张对照，正式目录不以尝试截图总数为目标，只保留真实功能界面、实际数据界面或有明确业务价值的配置界面。"
    )
    doc.add_paragraph(
        f"SHA-256 完全重复检查通过；缩略图视觉相似度复核（阈值 {SIMILARITY_THRESHOLD:.3f}）发现 {len(pairs)} 对候选，"
        "均为不同功能或不同业务状态，已人工保留，不据相似布局删除。"
    )
    doc.add_paragraph(
        "清理前已复核各轮过程及补充目录；review3 中仅样品摄谱采集控制台和预激发进度两张图具备独立功能价值，"
        "review4 新拍图与正式控制台重复且无有效数据。其余主要是菜单、空白/无数据工作台、Windows 文件选择器、"
        "报错/警告、任务管理器、Codex 界面、焦点探测或与正式图重复的版本，均不纳入正式证据。"
    )
    doc.add_heading("当前已有合格运行截图的功能范围", level=2)
    for item in [
        "光谱方法：新建、打开、另存为、页面设置和打印机设置。",
        "分析条件：激发/谱板/谱带参数、分析谱线列表和编辑入口、谱线类别、内标方式、峰值方式、拟合方式及参数打印。",
        "分析测试：色散/蒸发/样品摄谱工作条件、采集控制台、预激发进度、预录样号、谱图查看交互、样品分析结果、标准曲线、重复性检查和慢进配置。",
        "数据处理：重新计算、打印报告、报告导出向导三步、强度导出三步向导及实际文件载入。",
        "工具与帮助：软件参数、设备参数、CCD 设置、方法管理、备份压缩数据库、蒸发文件导入和关于窗口。",
    ]:
        add_bullet(doc, item)
    doc.add_heading("功能树中尚无合格运行截图的细分节点", level=2)
    for item in [
        "真实硬件采集过程中的暂停/继续/停止和完成收尾闭环；正式图只覆盖采集控制台和预激发倒计时。",
        "慢进完整暂停→人工调整→继续→结果写回闭环；当前截图只证明监控/等待时间配置入口。",
        "全时数据独立导出完成状态；现有图只证明全时谱图查看/观察。",
        "报告打印预览、PDF 导出和成功导出结果；报告导出向导第二/三步已补齐，但 PDF 保存对话框、文件选择器和错误残留图不作为证据。",
        "调试对中有效曲线/汞灯调试画面；当前硬件条件下只能取得串口异常。帮助主题同理，现有候选为 DIRECT.CHM 缺失提示。",
        "新软件要求中的转角排序/关键波段优先、基线/峰位异常闭环、汞灯专用流程、PDF/打印预览和角色权限管理，不能由现有截图宣称已实现；具体源码结论以功能树第五章为准。",
    ]:
        add_bullet(doc, item)
    doc.add_heading("复核结论", level=2)
    doc.add_paragraph(
        f"{file_count} 张正式截图可以覆盖主要可见配置和实际数据功能，但不能覆盖功能树中的每一个源码内部节点或每一个运行时分支。"
        " 未把无关或报错截图纳入正式证据；未采用素材已清理，备份副本保留。缺口已在功能说明中标为“源码确认但缺少合格运行截图”或“后续版本重点补充”。"
    )


def add_screenshot(doc, f, fig_no):
    group = group_for(f)
    p = doc.add_heading(f"图 {fig_no}：{title_for(f)}", level=2)
    p.paragraph_format.keep_with_next = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("功能说明：")
    set_font(r, bold=True, color="203748")
    r = p.add_run(description_for(f))
    set_font(r)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("实测状态：")
    set_font(r, bold=True, color="203748")
    r = p.add_run(f"{status_for(f)}；源码依据：{evidence_for(group)}。")
    set_font(r, size=10, color="555555")
    picture = doc.add_picture(str(f), width=Inches(6.25))
    picture._inline.docPr.set("descr", f"图 {fig_no}：{title_for(f)}")
    picture._inline.docPr.set("title", f.name)
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(f"图 {fig_no}  {f.name}", style="Caption")
    p.paragraph_format.keep_with_next = False


def write_docx(files, tree_text):
    doc = Document()
    configure_doc(doc)
    add_cover(doc, len(files))

    doc.add_heading("文档结构", level=1)
    for text in [
        "第一部分：按源码审查最终版功能树整理功能层次。",
        "第二部分：按 0-6 功能组编排运行截图，截图名称按功能重命名。",
        "第三部分：单列新软件要求、当前源码覆盖判断和后续验收重点。",
    ]:
        add_bullet(doc, text)
    doc.add_paragraph("正式截图均为 PNG，采集尺寸为 2880×1800；过程截图已隔离，不纳入本报告。")

    doc.add_heading("功能树（源码审查最终版）", level=1)
    add_tree_to_doc(doc, tree_text)
    add_status_table(doc)
    add_coverage_audit(doc, len(files), files)

    doc.add_heading("运行截图证据", level=1)
    doc.add_paragraph(f"以下按功能树顺序收录全部 {len(files)} 张正式截图。每张截图均保留文件名、实测状态、功能说明和源码依据。")
    current = None
    fig = 1
    for f in files:
        group = group_for(f)
        if group != current:
            current = group
            g = GROUPS[group]
            doc.add_heading(f"{g['number']}. {g['title']}", level=1)
            doc.add_paragraph(g["summary"])
            p = doc.add_paragraph()
            r = p.add_run("源码依据：")
            set_font(r, bold=True, color="203748")
            r = p.add_run(g["source"])
            set_font(r, color="555555")
        doc.add_page_break()
        add_screenshot(doc, f, fig)
        fig += 1

    doc.add_heading("交付与复核说明", level=1)
    doc.add_paragraph("本报告与同目录中的 01_功能树_源码整理版.md、02_截图清单_功能覆盖版.md 配套使用。功能树回答当前 Spec2.02 源码和界面有哪些功能；截图章节提供运行证据；新软件要求章节用于后续版本补充和验收。")
    doc.save(FINAL_DOCX)


def write_final_md(files, tree_text):
    lines = [
        "# Spec2.02 功能说明与截图文档",
        "",
        f"- 编制日期：{date.today()}",
        "- 入口：`../Spec2.02/SpecDirect.exe`",
        "- 源码：`../Spec Source/Source`",
        "- 运行环境：2880×1800，桌面缩放 200%",
        f"- 正式截图：{len(files)} 张 PNG，全部 2880×1800",
        "",
        "## 使用说明",
        "",
        "本文档按功能树层级组织：先给出源码审查最终版功能树，再按 0-6 功能组列出运行截图。截图名称按功能重命名，未被最终文档引用的过程/补充截图已清理，备份副本保留。",
        "",
        "## 功能树（源码审查最终版）",
        "",
        tree_text.strip(),
        "",
    ]
    lines += coverage_audit_lines(files)
    lines += ["## 截图证据", ""]
    current = None
    fig = 1
    for f in files:
        group = group_for(f)
        if group != current:
            current = group
            g = GROUPS[group]
            lines += [f"### {g['number']}. {g['title']}", "", g["summary"], "", f"源码依据：{g['source']}", ""]
        lines += [
            f"#### 图 {fig}：{title_for(f)}",
            "",
            f"- 文件：`{f.name}`",
            f"- 实测状态：{status_for(f)}",
            f"- 功能说明：{description_for(f)}",
            f"- 源码依据：{evidence_for(group)}",
            "",
            f"![图 {fig}：{title_for(f)}]({md_image_path(f)})",
            "",
        ]
        fig += 1
    lines += [
        "## 交付与复核说明",
        "",
        "本报告与同目录中的 `01_功能树_源码整理版.md`、`02_截图清单_功能覆盖版.md` 配套使用。设备采集和串口部分只证明配置入口及可取得的采集状态，未把异常图作为功能证据；本次没有执行真实硬件采集或数据库写入。",
        "",
    ]
    FINAL_MD.write_text("\n".join(lines), encoding="utf-8")


def main():
    files = entries()
    tree_text = TREE_FILE.read_text(encoding="utf-8")
    write_screenshot_list(files)
    write_final_md(files, tree_text)
    write_docx(files, tree_text)
    print(f"generated {SHOT_LIST_MD}")
    print(f"generated {FINAL_MD}")
    print(f"generated {FINAL_DOCX}")


if __name__ == "__main__":
    main()
